import os
from typing import List

try:
    import PyPDF2
    PDF_PARSER = "PyPDF2"
except ImportError:
    PDF_PARSER = None

try:
    import pdfplumber
    PDF_PARSER = "pdfplumber"
except ImportError:
    pass

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_text_from_pdf(file_path: str) -> str:
    """从 PDF 提取文本"""
    text = ""

    if PDF_PARSER == "pdfplumber":
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    else:
        import PyPDF2
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

    return text


def extract_text_from_docx(file_path: str) -> str:
    """从 Word 文档提取文本"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 未安装")

    doc = DocxDocument(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"

    return text


def extract_text(file_path: str) -> str:
    """根据文件扩展名提取文本"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 200) -> List[str]:
    """将长文本分块"""
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]

        if end < text_length:
            last_punct = max(
                chunk.rfind("。"),
                chunk.rfind("？"),
                chunk.rfind("\n"),
                chunk.rfind(".")
            )
            if last_punct > chunk_size - 200:
                end = start + last_punct + 1
                chunk = text[start:end]

        chunks.append(chunk)
        start = end - overlap

    return chunks


def chunk_document(file_path: str, chunk_size: int = 2000, chunk_overlap: int = 200) -> List[dict]:
    """对文档进行分块"""
    text = extract_text(file_path)
    chunks = chunk_text(text, chunk_size, chunk_overlap)
    file_type = os.path.splitext(file_path)[1].lower()

    return [
        {
            "content": chunk,
            "source": file_path,
            "file_type": file_type
        }
        for chunk in chunks
    ]
